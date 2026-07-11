# -*- coding: utf-8 -*-
"""掛籠資訊表產生器（可編輯 Word）。

由外部資料檔產生「貓咪掛籠資訊表」：上方貓咪橫幅 + 9 欄可編輯表格 + 下方貓咪腳印，
每張 81.1×79mm、每頁 6 張（2×3）。文字皆可於 Word 直接編輯，對位由表格保證。

用法：
    python generate_cage_cards.py --pets pets.json --licenses licenses.json --out 掛籠資訊表.docx

資料檔格式見 pets.sample.json / licenses.sample.json。
真實個資（晶片號、證號、貓舍名）請放在自己的 pets.json / licenses.json，
這些檔案已被 .gitignore 排除，不會進版控。
"""
import argparse
import json
import os

from docx import Document
from docx.shared import Mm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")

# 版面常數（mm）—— 對應原範本量測尺寸
CARD_W = 81.1
BANNER_H, FOOTER_H, ROW_H = 13.4, 15.0, 5.62   # 合計 ≈ 79mm
LAB_W = 30.0
COLS, ROWS = 2, 3          # 每頁 2×3 = 6 張
PINK, BG, BORDER = "FBE2EA", "FBF3D9", "E3B7C6"
INK = RGBColor(0x3A, 0x2F, 0x2A)
LABINK = RGBColor(0x7A, 0x4A, 0x5A)
FIELDS = ["寵物類別", "品種", "性別", "生日", "晶片號碼", "絕育情形",
          "毛色特徵", "來源許可證號", "來源母貓晶片號碼"]
SMALL_FIELDS = {"晶片號碼", "來源許可證號", "來源母貓晶片號碼"}


def fmt_chip(c):
    c = "".join(ch for ch in str(c) if ch.isdigit())
    return "-".join([c[0:5], c[5:10], c[10:15]]) if len(c) == 15 else c


# ---- 低階 OOXML 輔助 ----
def _rowh(row, mm, exact=True):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight'); h.set(qn('w:val'), str(int(mm * 56.7)))
    h.set(qn('w:hRule'), 'exact' if exact else 'atLeast'); trPr.append(h)
    trPr.append(OxmlElement('w:cantSplit'))


def _cellw(cell, mm):
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(mm * 56.7))); w.set(qn('w:type'), 'dxa')
    cell._tc.get_or_add_tcPr().append(w)


def _shade(cell, hexc):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexc)
    cell._tc.get_or_add_tcPr().append(sh)


def _borders(tbl, color, sz, style='single'):
    b = OxmlElement('w:tblBorders')
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        x = OxmlElement('w:' + e); x.set(qn('w:val'), style)
        if style == 'single':
            x.set(qn('w:sz'), str(sz)); x.set(qn('w:space'), '0'); x.set(qn('w:color'), color)
        b.append(x)
    tbl._tbl.tblPr.append(b)


def _margins(tbl, v=0, h=40):
    mar = OxmlElement('w:tblCellMar')
    for t, val in (('top', v), ('bottom', v), ('start', h), ('end', h)):
        x = OxmlElement('w:' + t); x.set(qn('w:w'), str(val)); x.set(qn('w:type'), 'dxa'); mar.append(x)
    tbl._tbl.tblPr.append(mar)


def _img(cell, path, w_mm, h_mm):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(path, width=Emu(int(w_mm * 36000)), height=Emu(int(h_mm * 36000)))


def _text(cell, text, *, size, bold, align, color):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'c' else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = color
    r.font.name = 'Microsoft JhengHei'
    r._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')


def build_card(cell, pet, licenses):
    station = pet.get("station", "")
    vals = [
        "貓", pet.get("breed", ""), pet.get("sex", ""), pet.get("birth", ""),
        fmt_chip(pet.get("chip", "")), pet.get("neuter", "未絕育"),
        pet.get("color", ""), licenses.get(station, pet.get("license", "")),
        fmt_chip(pet.get("dam_chip", "")),
    ]
    t = cell.add_table(rows=len(FIELDS) + 2, cols=2)
    _borders(t, BORDER, 4); _margins(t)
    banner = t.cell(0, 0).merge(t.cell(0, 1)); _shade(banner, BG)
    _rowh(t.rows[0], BANNER_H); _img(banner, os.path.join(ASSETS, "banner.png"), CARD_W, BANNER_H)
    for i, (f, v) in enumerate(zip(FIELDS, vals), start=1):
        _rowh(t.rows[i], ROW_H)
        lc, vc = t.cell(i, 0), t.cell(i, 1)
        _shade(lc, PINK); _cellw(lc, LAB_W); _cellw(vc, CARD_W - LAB_W)
        _text(lc, f, size=7.5, bold=True, align='c', color=LABINK)
        _text(vc, str(v), size=7.5 if f in SMALL_FIELDS else 9, bold=True, align='l', color=INK)
    footer = t.cell(len(FIELDS) + 1, 0).merge(t.cell(len(FIELDS) + 1, 1)); _shade(footer, BG)
    _rowh(t.rows[len(FIELDS) + 1], FOOTER_H)
    _img(footer, os.path.join(ASSETS, "footer.png"), CARD_W, FOOTER_H)


def build(pets, licenses, out_path):
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Mm(7))
    st = doc.styles['Normal']
    st.paragraph_format.space_after = Pt(0); st.paragraph_format.line_spacing = 1.0

    per = COLS * ROWS
    for gi in range(0, len(pets), per):
        group = pets[gi:gi + per]
        outer = doc.add_table(rows=ROWS, cols=COLS)
        outer.alignment = WD_TABLE_ALIGNMENT.CENTER
        _borders(outer, "", 0, style='none'); _margins(outer, 0, 60)
        for r in outer.rows:
            _rowh(r, 80.5, exact=False)
        for j, pet in enumerate(group):
            oc = outer.cell(j // COLS, j % COLS); _cellw(oc, CARD_W + 2)
            build_card(oc, pet, licenses)
        if gi + per < len(pets):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.save(out_path)
    return (len(pets) + per - 1) // per


def main():
    ap = argparse.ArgumentParser(description="產生可編輯掛籠資訊表 Word 檔")
    ap.add_argument("--pets", default=os.path.join(HERE, "pets.json"), help="寵物資料 JSON")
    ap.add_argument("--licenses", default=os.path.join(HERE, "licenses.json"), help="登記站→特寵業字號 對照 JSON")
    ap.add_argument("--out", default=os.path.join(HERE, "掛籠資訊表.docx"), help="輸出 docx 路徑")
    args = ap.parse_args()

    with open(args.pets, encoding="utf-8") as f:
        pets = json.load(f)
    licenses = {}
    if os.path.exists(args.licenses):
        with open(args.licenses, encoding="utf-8") as f:
            licenses = json.load(f)

    pages = build(pets, licenses, args.out)
    print(f"已產生 {args.out}：{len(pets)} 張、{pages} 頁")


if __name__ == "__main__":
    main()
